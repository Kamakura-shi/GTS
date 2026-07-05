#!/usr/bin/env python3
"""Genere un questionnaire (GTS) au format .docx.

Contenu :
  - Partie A : 4 questions de developpement (reponse en 1-2 phrases)
  - Partie B : 20 QCM DICOM (8 choix, 4 reponses correctes)
  - Partie C : 10 QCM HL7v2 (8 choix, 4 reponses correctes)
  - Corrige : reponses modeles + choix corrects
"""

import random

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Partie A : questions de developpement (reponse courte 1-2 phrases)
# ---------------------------------------------------------------------------
DEV_QUESTIONS = [
    {
        "q": "Expliquez en une ou deux phrases la difference fondamentale entre "
             "DICOM et HL7v2 ainsi que leurs domaines d'application respectifs.",
        "a": "DICOM est le standard d'echange et de stockage des images medicales "
             "(et de leurs metadonnees) ainsi que du protocole reseau associe, "
             "tandis que HL7v2 est un standard de messagerie pour l'echange de "
             "donnees administratives et cliniques (admissions, demandes, "
             "resultats) entre systemes d'information de sante.",
    },
    {
        "q": "Decrivez brievement le modele de donnees hierarchique de DICOM "
             "(Patient, Etude, Serie, Instance).",
        "a": "DICOM organise les donnees de facon hierarchique : un Patient possede "
             "une ou plusieurs Etudes, chaque Etude regroupe une ou plusieurs Series "
             "(en general une acquisition d'une meme modalite), et chaque Serie "
             "contient une ou plusieurs Instances (images), chacune identifiee par "
             "un SOP Instance UID unique.",
    },
    {
        "q": "Qu'est-ce qu'un PACS et quel role joue-t-il dans un service d'imagerie "
             "medicale ?",
        "a": "Un PACS (Picture Archiving and Communication System) est le systeme qui "
             "archive, gere et distribue les images DICOM ; il communique avec les "
             "modalites et les stations de visualisation via le protocole DICOM pour "
             "centraliser le stockage et la consultation des examens.",
    },
    {
        "q": "Expliquez le role du segment MSH dans un message HL7v2.",
        "a": "Le segment MSH (Message Header) est le segment obligatoire qui debute "
             "tout message HL7v2 : il definit les caracteres de separation, identifie "
             "les applications/etablissements emetteur et recepteur, et precise le "
             "type de message et l'evenement declencheur.",
    },
]

# ---------------------------------------------------------------------------
# Chaque QCM : (enonce, [8 propositions], {indices des 4 propositions vraies})
# ---------------------------------------------------------------------------
DICOM_MCQS = [
    ("Concernant les generalites de DICOM, quelles propositions sont exactes ?",
     ["DICOM signifie \"Digital Imaging and Communications in Medicine\".",
      "DICOM definit a la fois un format de fichier et un protocole de communication reseau.",
      "Le standard DICOM est maintenu par le NEMA (comite DICOM).",
      "Un fichier DICOM regroupe dans un meme objet les metadonnees (patient, examen) et les donnees image.",
      "DICOM est un standard exclusivement destine au stockage local, sans dimension reseau.",
      "DICOM ne peut coder que des images en niveaux de gris.",
      "DICOM a ete concu et est maintenu par l'OMS.",
      "Un fichier DICOM ne contient aucune information relative au patient."],
     {0, 1, 2, 3}),

    ("Concernant la structure d'un fichier DICOM, quelles propositions sont exactes ?",
     ["Un fichier DICOM est compose d'une suite d'elements de donnees (Data Elements).",
      "Chaque Data Element est identifie par un tag (numero de groupe + numero d'element).",
      "Le preambule d'un fichier DICOM fait 128 octets, suivi du prefixe \"DICM\".",
      "Le Data Set contient les attributs decrivant l'objet (IOD).",
      "Les tags DICOM sont exprimes uniquement en base 10.",
      "Le preambule DICOM fait 512 octets.",
      "Un fichier DICOM ne peut contenir qu'un seul Data Element.",
      "Le prefixe magique d'un fichier DICOM est \"HL7\"."],
     {0, 1, 2, 3}),

    ("Concernant les tags DICOM, quelles propositions sont exactes ?",
     ["Un tag DICOM s'ecrit sous la forme (gggg,eeee) en hexadecimal.",
      "Le groupe (0010) est associe aux informations relatives au patient.",
      "Le groupe (0008) contient des informations generales sur l'etude/l'image.",
      "Les groupes impairs sont reserves aux elements prives (private tags).",
      "Les tags de groupe pair sont reserves aux fabricants.",
      "Le tag (7FE0,0010) contient le nom du patient.",
      "Un tag est code sur 8 octets de texte ASCII lisibles.",
      "Le numero de groupe et le numero d'element sont identiques pour tous les objets."],
     {0, 1, 2, 3}),

    ("Concernant la Value Representation (VR), quelles propositions sont exactes ?",
     ["La VR decrit le type de donnees de la valeur d'un element.",
      "\"PN\" est la VR utilisee pour un nom de personne (Person Name).",
      "\"DA\" correspond a une date.",
      "\"UI\" correspond a un identifiant unique (UID).",
      "La VR \"PN\" designe un nombre a virgule.",
      "Il n'existe qu'une seule VR possible dans DICOM.",
      "La VR \"US\" signifie \"ultrason\".",
      "La VR n'a aucune influence sur l'encodage de la valeur."],
     {0, 1, 2, 3}),

    ("Concernant la Transfer Syntax, quelles propositions sont exactes ?",
     ["La Transfer Syntax definit l'encodage des donnees (ordre des octets, VR explicite/implicite, compression).",
      "Il existe une Transfer Syntax \"Implicit VR Little Endian\" (par defaut).",
      "Certaines Transfer Syntax permettent la compression JPEG des images.",
      "Une Transfer Syntax est identifiee par un UID.",
      "La Transfer Syntax decrit l'identite du patient.",
      "La Transfer Syntax est toujours stockee au format XML en clair.",
      "La compression est interdite en DICOM.",
      "La Transfer Syntax fixe la date de l'examen."],
     {0, 1, 2, 3}),

    ("Concernant le modele de donnees hierarchique de DICOM, quelles propositions sont exactes ?",
     ["Le modele est hierarchique : Patient > Etude > Serie > Instance.",
      "Une etude peut contenir plusieurs series.",
      "Une serie correspond generalement a une acquisition d'une meme modalite.",
      "Chaque instance possede un SOP Instance UID unique.",
      "Une serie ne peut contenir qu'une seule image.",
      "Le patient se situe au bas de la hierarchie, sous l'instance.",
      "Un meme Study Instance UID est partage par tous les patients.",
      "Il n'existe pas de notion de serie en DICOM."],
     {0, 1, 2, 3}),

    ("Concernant les notions de SOP et d'IOD, quelles propositions sont exactes ?",
     ["SOP signifie Service-Object Pair.",
      "Une SOP Class associe un IOD a un ensemble de services (DIMSE).",
      "Un IOD (Information Object Definition) definit les attributs d'un type d'objet.",
      "Chaque SOP Class est identifiee par un SOP Class UID.",
      "IOD signifie \"Image Only Data\".",
      "Une SOP Class ne peut pas etre identifiee par un UID.",
      "Le SOP Instance UID est identique pour toutes les images d'une etude.",
      "Un IOD ne decrit que la couleur des pixels."],
     {0, 1, 2, 3}),

    ("Concernant les services DIMSE, quelles propositions sont exactes ?",
     ["C-STORE permet d'envoyer/stocker un objet sur un noeud distant.",
      "C-FIND permet d'interroger (query) une base d'objets.",
      "C-MOVE permet de demander le transfert d'objets vers une destination.",
      "C-ECHO permet de verifier la connectivite (ping applicatif).",
      "C-STORE sert a supprimer definitivement une image.",
      "C-ECHO transfere les images compressees.",
      "C-FIND cree un nouveau patient dans le RIS.",
      "C-MOVE sert uniquement a imprimer un film."],
     {0, 1, 2, 3}),

    ("Concernant les AE Titles et l'association DICOM, quelles propositions sont exactes ?",
     ["Une Application Entity (AE) est identifiee par un AE Title.",
      "L'etablissement d'une communication DICOM passe par une negociation d'association.",
      "La negociation inclut les presentation contexts (SOP Class + Transfer Syntax).",
      "L'AE Title, l'adresse IP et le port sont necessaires pour joindre un noeud DICOM.",
      "L'AE Title est toujours l'adresse MAC de la machine.",
      "Aucune negociation n'est necessaire avant l'echange d'objets.",
      "Le port DICOM est obligatoirement le port 80.",
      "Un presentation context ne contient que l'adresse IP."],
     {0, 1, 2, 3}),

    ("Concernant le PACS et la Modality Worklist, quelles propositions sont exactes ?",
     ["Un PACS archive et distribue les images DICOM.",
      "La Modality Worklist fournit a la modalite la liste des examens programmes.",
      "La Modality Worklist s'appuie sur le service C-FIND.",
      "Le PACS communique avec les modalites via le protocole DICOM.",
      "Un PACS ne peut stocker que des documents PDF.",
      "La Modality Worklist impose de saisir manuellement le patient sur chaque machine.",
      "Le PACS remplace le protocole DICOM par du FTP simple.",
      "La Worklist empeche toute communication reseau."],
     {0, 1, 2, 3}),

    ("Concernant les donnees image (Pixel Data), quelles propositions sont exactes ?",
     ["Les donnees image sont stockees dans l'element (7FE0,0010) Pixel Data.",
      "L'attribut Photometric Interpretation indique par ex. MONOCHROME2 ou RGB.",
      "Rows et Columns decrivent les dimensions de l'image.",
      "Bits Allocated indique le nombre de bits alloues par echantillon.",
      "DICOM ne stocke jamais les pixels de l'image.",
      "MONOCHROME2 signifie que l'image est forcement en couleur RVB.",
      "Rows et Columns decrivent la date de naissance du patient.",
      "Le nombre de bits par pixel est toujours de 1."],
     {0, 1, 2, 3}),

    ("Concernant les UID en DICOM, quelles propositions sont exactes ?",
     ["Un UID garantit l'unicite globale d'un objet ou d'un concept.",
      "Les UID sont composes de chiffres separes par des points.",
      "Study/Series/SOP Instance UID sont des exemples d'UID.",
      "Un UID commence souvent par une racine attribuee a l'organisation (org root).",
      "Un UID est une chaine de texte libre non structuree.",
      "Deux objets differents doivent partager le meme UID.",
      "Les UID contiennent obligatoirement des lettres accentuees.",
      "Un UID change a chaque ouverture du fichier."],
     {0, 1, 2, 3}),

    ("Concernant la compression en DICOM, quelles propositions sont exactes ?",
     ["DICOM supporte la compression JPEG (avec et sans perte).",
      "DICOM supporte JPEG 2000.",
      "DICOM supporte la compression RLE (Run Length Encoding).",
      "Le type de compression est indique par la Transfer Syntax.",
      "DICOM interdit toute forme de compression.",
      "La compression sans perte modifie irreversiblement les pixels.",
      "La compression est indiquee dans le nom du patient.",
      "RLE signifie \"Radiology Long Exam\"."],
     {0, 1, 2, 3}),

    ("Concernant les modalites, quelles propositions sont exactes ?",
     ["CT designe la tomodensitometrie (scanner).",
      "MR designe l'IRM (resonance magnetique).",
      "US designe l'echographie (ultrasons).",
      "CR/DX designent la radiographie numerique.",
      "MR designe la mammographie analogique sur film.",
      "US designe un examen de medecine nucleaire par positons.",
      "CT designe un electrocardiogramme.",
      "Toutes les modalites partagent le meme code \"XX\"."],
     {0, 1, 2, 3}),

    ("Concernant le DICOMDIR, quelles propositions sont exactes ?",
     ["Un fichier DICOMDIR sert d'index des objets DICOM presents sur un support (CD/DVD).",
      "DICOMDIR decrit la hierarchie Patient/Study/Series/Image du media.",
      "DICOMDIR facilite la navigation sans lire tous les fichiers.",
      "Un CD patient de radiologie contient souvent un DICOMDIR a la racine.",
      "DICOMDIR est un format d'image compresse.",
      "DICOMDIR contient les pixels de toutes les images en double.",
      "DICOMDIR est un protocole reseau remplacant C-STORE.",
      "DICOMDIR interdit le stockage sur support amovible."],
     {0, 1, 2, 3}),

    ("Concernant le DICOM Conformance Statement, quelles propositions sont exactes ?",
     ["Il decrit les fonctionnalites DICOM supportees par un produit.",
      "Il precise les SOP Classes supportees (SCU/SCP).",
      "Il indique les Transfer Syntaxes supportees.",
      "Il est utile pour evaluer l'interoperabilite entre systemes.",
      "Il contient le dossier medical complet du patient.",
      "Il garantit a lui seul une interoperabilite parfaite sans test.",
      "SCU/SCP y designent des formats d'image.",
      "Le Conformance Statement est interdit par le standard."],
     {0, 1, 2, 3}),

    ("Concernant les elements d'identification patient/etude, quelles propositions sont exactes ?",
     ["Patient Name (0010,0010) contient le nom du patient.",
      "Patient ID (0010,0020) contient l'identifiant du patient.",
      "Study Date (0008,0020) contient la date de l'etude.",
      "Modality (0008,0060) contient le type de modalite.",
      "Patient Name est stocke dans le tag (7FE0,0010).",
      "Modality contient la date de naissance du patient.",
      "Patient ID contient les pixels de l'image.",
      "Study Date est stockee dans le groupe (0010)."],
     {0, 1, 2, 3}),

    ("Concernant les roles SCU et SCP, quelles propositions sont exactes ?",
     ["Le SCU (Service Class User) est le \"client\" qui demande un service.",
      "Le SCP (Service Class Provider) est le \"serveur\" qui fournit le service.",
      "Lors d'un C-STORE, la modalite agit typiquement en SCU et le PACS en SCP.",
      "Un meme equipement peut etre SCU pour un service et SCP pour un autre.",
      "SCU et SCP designent deux formats de compression.",
      "Le PACS ne peut jamais etre SCP.",
      "Le role SCU/SCP n'existe pas en DICOM.",
      "SCU signifie \"Standard Compression Unit\"."],
     {0, 1, 2, 3}),

    ("Concernant d'autres Value Representations, quelles propositions sont exactes ?",
     ["\"DT\" represente une date/heure (DateTime).",
      "\"TM\" represente une heure (Time).",
      "\"IS\" represente un entier code en chaine (Integer String).",
      "\"DS\" represente un decimal code en chaine (Decimal String).",
      "\"TM\" represente une image miniature.",
      "\"IS\" represente obligatoirement le sexe du patient.",
      "\"DS\" represente un son.",
      "Aucune de ces VR n'existe dans DICOM."],
     {0, 1, 2, 3}),

    ("Concernant l'interoperabilite et l'integration de DICOM, quelles propositions sont exactes ?",
     ["DICOM est largement utilise en imagerie medicale (radiologie, cardiologie...).",
      "DICOM peut etre integre avec HL7 dans un systeme d'information hospitalier.",
      "IHE propose des profils d'integration s'appuyant sur DICOM et HL7.",
      "DICOM permet l'echange d'images entre equipements de fabricants differents.",
      "DICOM est incompatible avec tout autre standard de sante.",
      "DICOM ne fonctionne qu'avec un seul fabricant proprietaire.",
      "DICOM remplace totalement le dossier patient administratif.",
      "DICOM interdit toute interoperabilite multi-constructeurs."],
     {0, 1, 2, 3}),
]

HL7_MCQS = [
    ("Concernant les generalites de HL7v2, quelles propositions sont exactes ?",
     ["HL7 signifie \"Health Level Seven\".",
      "HL7v2 est un standard d'echange de messages entre systemes d'information de sante.",
      "Le nom fait reference a la couche 7 (application) du modele OSI.",
      "HL7v2 est tres repandu pour les echanges administratifs et cliniques (ADT, resultats...).",
      "HL7 est un format d'image medicale concurrent de DICOM.",
      "HL7 signifie \"Hospital Level 7 imaging\".",
      "HL7v2 code uniquement des images de scanner.",
      "HL7 n'a aucun lien avec les systemes d'information de sante."],
     {0, 1, 2, 3}),

    ("Concernant la structure d'un message HL7v2, quelles propositions sont exactes ?",
     ["Un message HL7v2 est compose de segments.",
      "Chaque segment est compose de champs (fields).",
      "Les champs peuvent contenir des composants et sous-composants.",
      "Chaque segment commence par un code a 3 lettres (ex. MSH, PID).",
      "Un message HL7v2 est un fichier binaire d'image.",
      "Les segments HL7 n'ont pas de nom.",
      "Un message HL7v2 ne contient qu'un seul caractere.",
      "Les champs HL7 ne peuvent jamais contenir de composants."],
     {0, 1, 2, 3}),

    ("Concernant les delimiteurs HL7v2, quelles propositions sont exactes ?",
     ["Le caractere \"|\" (pipe) separe les champs.",
      "Le caractere \"^\" separe les composants.",
      "Le caractere \"~\" separe les repetitions d'un champ.",
      "Le caractere \"&\" separe les sous-composants.",
      "Le \"|\" separe les sous-composants.",
      "Les delimiteurs HL7 sont interdits et n'existent pas.",
      "Le \"^\" separe les segments entre eux.",
      "Les composants d'un champ sont separes par un retour a la ligne."],
     {0, 1, 2, 3}),

    ("Concernant le segment MSH, quelles propositions sont exactes ?",
     ["Le segment MSH (Message Header) est obligatoire et debute le message.",
      "MSH contient les applications/etablissements emetteur et recepteur.",
      "MSH precise le type de message et l'evenement declencheur.",
      "MSH definit les caracteres de separation utilises dans le message.",
      "MSH est un segment optionnel place a la fin du message.",
      "MSH contient les pixels d'une image.",
      "MSH n'apparait jamais dans un message HL7v2.",
      "MSH sert a stocker le resultat d'une IRM au format DICOM."],
     {0, 1, 2, 3}),

    ("Concernant les types de messages HL7v2, quelles propositions sont exactes ?",
     ["ADT concerne les mouvements/admissions de patients (Admit, Discharge, Transfer).",
      "ORM concerne les demandes/commandes (orders).",
      "ORU concerne la transmission de resultats d'observations.",
      "SIU concerne la gestion des rendez-vous/planification.",
      "ADT est un format de compression d'image.",
      "ORU sert a eteindre le serveur.",
      "ORM concerne uniquement l'impression de films radiologiques.",
      "Les types de message HL7 n'existent pas."],
     {0, 1, 2, 3}),

    ("Concernant les evenements declencheurs (trigger events), quelles propositions sont exactes ?",
     ["A01 correspond a l'admission d'un patient (Admit).",
      "A02 correspond a un transfert de patient.",
      "A03 correspond a la sortie (Discharge) d'un patient.",
      "Le trigger event est associe au type de message (ex. ADT^A01).",
      "A01 correspond a l'envoi d'une image CT.",
      "Le trigger event n'a aucun rapport avec le type de message.",
      "A03 signifie \"ajouter 3 patients\".",
      "Les evenements declencheurs sont reserves a DICOM."],
     {0, 1, 2, 3}),

    ("Concernant les segments courants HL7v2, quelles propositions sont exactes ?",
     ["PID contient les donnees d'identification du patient.",
      "PV1 contient les donnees de la venue/visite (Patient Visit).",
      "OBR decrit une demande/observation (Observation Request).",
      "OBX contient une observation/resultat elementaire (Observation/Result).",
      "PID contient les parametres reseau DICOM.",
      "OBX est le segment d'en-tete obligatoire du message.",
      "PV1 contient les pixels de l'image.",
      "OBR signifie \"Object Binary Raw\" pour les images."],
     {0, 1, 2, 3}),

    ("Concernant l'accuse de reception (ACK) HL7v2, quelles propositions sont exactes ?",
     ["HL7v2 prevoit un message d'accuse de reception (ACK).",
      "Le segment MSA (Message Acknowledgment) indique le statut de traitement.",
      "Un code AA signifie une acceptation (Application Accept).",
      "Un code AE ou AR indique une erreur/refus applicatif.",
      "HL7v2 ne dispose d'aucun mecanisme d'acquittement.",
      "Le segment MSA contient les donnees image.",
      "AA signifie \"Absolute Anonymization\".",
      "L'accuse de reception est interdit par le standard."],
     {0, 1, 2, 3}),

    ("Concernant l'encodage des messages HL7v2, quelles propositions sont exactes ?",
     ["L'encodage \"pipe and hat\" (ER7) est l'encodage historique de HL7v2.",
      "HL7v2 dispose aussi d'une representation XML.",
      "Un message ER7 est lisible sous forme texte avec des delimiteurs.",
      "Les segments d'un message ER7 sont separes par un retour chariot (CR).",
      "HL7v2 est uniquement un format binaire chiffre illisible.",
      "L'encodage ER7 interdit tout delimiteur.",
      "HL7v2 ne peut jamais etre represente en XML.",
      "Les segments ER7 sont obligatoirement separes par des virgules."],
     {0, 1, 2, 3}),

    ("Concernant HL7v2 par rapport a HL7v3 et FHIR, quelles propositions sont exactes ?",
     ["HL7v2 est base sur des messages et des delimiteurs, tres deploye en pratique.",
      "HL7 FHIR est un standard plus recent base sur des ressources et des API REST.",
      "HL7v3 s'appuie sur un modele d'information (RIM) et le XML.",
      "HL7v2 et FHIR peuvent coexister dans un meme etablissement.",
      "HL7v2 et FHIR sont exactement le meme standard avec le meme encodage.",
      "FHIR est un format d'image DICOM.",
      "HL7v3 est anterieur a HL7v2.",
      "HL7v2 n'est plus utilise nulle part aujourd'hui."],
     {0, 1, 2, 3}),
]

LETTERS = "abcdefgh"


def shuffle_mcqs(mcqs, seed):
    """Melange les propositions de chaque QCM (deterministe) pour que les 4
    reponses correctes ne soient pas toujours a, b, c, d. Retourne une liste
    (enonce, propositions_melangees, {indices corrects}).
    """
    rng = random.Random(seed)
    prepared = []
    for q, opts, correct in mcqs:
        order = list(range(len(opts)))
        rng.shuffle(order)
        new_opts = [opts[i] for i in order]
        new_correct = {pos for pos, orig in enumerate(order) if orig in correct}
        prepared.append((q, new_opts, new_correct))
    return prepared


def add_heading(doc, text, size, color=(0, 0, 0), space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.space_before = Pt(space_before)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_mcq(doc, number, question, options):
    p = doc.add_paragraph()
    run = p.add_run(f"Q{number}. {question}")
    run.bold = True
    p.paragraph_format.space_before = Pt(10)
    hint = doc.add_paragraph()
    hr = hint.add_run("(QCM - 8 propositions, 4 reponses correctes)")
    hr.italic = True
    hr.font.size = Pt(9)
    hint.paragraph_format.space_after = Pt(2)
    for i, opt in enumerate(options):
        op = doc.add_paragraph(style="List Paragraph")
        op.paragraph_format.left_indent = Pt(24)
        op.paragraph_format.space_after = Pt(0)
        op.add_run(f"{LETTERS[i]})  {opt}")


def build():
    dicom = shuffle_mcqs(DICOM_MCQS, seed=20260704)
    hl7 = shuffle_mcqs(HL7_MCQS, seed=70260704)

    doc = Document()

    # ---- Page de titre ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Questionnaire - GTS")
    tr.bold = True
    tr.font.size = Pt(22)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Standards en sante : DICOM et HL7v2")
    sr.font.size = Pt(13)
    sr.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Nom : _______________________     Date : ____________\n")
    meta.add_run("Bareme : Partie A (4 x 3 pts) + Partie B (20 QCM) + Partie C (10 QCM)")

    note = doc.add_paragraph()
    nr = note.add_run(
        "Consignes : pour les QCM, chaque question comporte 8 propositions dont "
        "exactement 4 sont correctes ; cochez les 4 propositions vraies. Pour la "
        "partie A, repondez en une ou deux phrases.")
    nr.font.size = Pt(10)
    nr.italic = True

    # ---- Partie A ----
    add_heading(doc, "Partie A - Questions de developpement", 15, (0x1F, 0x49, 0x7D))
    for i, item in enumerate(DEV_QUESTIONS, start=1):
        p = doc.add_paragraph()
        r = p.add_run(f"A{i}. {item['q']}")
        r.bold = True
        p.paragraph_format.space_before = Pt(8)
        doc.add_paragraph("Reponse : ______________________________________________________")
        doc.add_paragraph("_______________________________________________________________")

    # ---- Partie B : DICOM ----
    add_heading(doc, "Partie B - DICOM (20 questions)", 15, (0x1F, 0x49, 0x7D))
    for i, (q, opts, _) in enumerate(dicom, start=1):
        add_mcq(doc, i, q, opts)

    # ---- Partie C : HL7v2 ----
    add_heading(doc, "Partie C - HL7v2 (10 questions)", 15, (0x1F, 0x49, 0x7D))
    for i, (q, opts, _) in enumerate(hl7, start=1):
        add_mcq(doc, i, q, opts)

    # ---- Corrige ----
    doc.add_page_break()
    add_heading(doc, "Corrige", 17, (0x8B, 0x00, 0x00))

    add_heading(doc, "Partie A - Elements de reponse attendus", 13, (0x8B, 0x00, 0x00))
    for i, item in enumerate(DEV_QUESTIONS, start=1):
        p = doc.add_paragraph()
        p.add_run(f"A{i}. ").bold = True
        p.add_run(item["a"])
        p.paragraph_format.space_after = Pt(6)

    add_heading(doc, "Partie B - DICOM : reponses correctes", 13, (0x8B, 0x00, 0x00))
    for i, (q, opts, correct) in enumerate(dicom, start=1):
        letters = ", ".join(LETTERS[j] for j in sorted(correct))
        p = doc.add_paragraph()
        p.add_run(f"Q{i} : ").bold = True
        p.add_run(letters + "  -  " + " / ".join(opts[j] for j in sorted(correct)))
        p.paragraph_format.space_after = Pt(4)

    add_heading(doc, "Partie C - HL7v2 : reponses correctes", 13, (0x8B, 0x00, 0x00))
    for i, (q, opts, correct) in enumerate(hl7, start=1):
        letters = ", ".join(LETTERS[j] for j in sorted(correct))
        p = doc.add_paragraph()
        p.add_run(f"Q{i} : ").bold = True
        p.add_run(letters + "  -  " + " / ".join(opts[j] for j in sorted(correct)))
        p.paragraph_format.space_after = Pt(4)

    import os
    out_dir = "Cr"
    os.makedirs(out_dir, exist_ok=True)
    out = os.path.join(out_dir, "Questionnaire_GTS_DICOM_HL7_Cr.docx")
    doc.save(out)
    # Sanity checks
    assert all(len(o) == 8 for _, o, _ in DICOM_MCQS + HL7_MCQS), "8 choix requis"
    assert all(len(c) == 4 for _, _, c in DICOM_MCQS + HL7_MCQS), "4 correctes requises"
    assert len(DEV_QUESTIONS) == 4
    assert len(DICOM_MCQS) == 20
    assert len(HL7_MCQS) == 10
    print(f"OK -> {out}")
    print(f"  Dev: {len(DEV_QUESTIONS)} | DICOM: {len(DICOM_MCQS)} | HL7v2: {len(HL7_MCQS)}")


if __name__ == "__main__":
    build()
